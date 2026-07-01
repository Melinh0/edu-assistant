import os
import tempfile
import subprocess
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT

def save_uploaded_file(uploaded_file, save_dir):
    file_path = os.path.join(save_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path

def get_files_list(directory):
    return [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]

def delete_file(filename, directory):
    os.remove(os.path.join(directory, filename))

def pandoc_convert(text, output_format, extra_args=None):
    if extra_args is None:
        extra_args = []
    engines = ['xelatex', 'pdflatex']  
    for engine in engines:
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as md_file:
                md_path = md_file.name
                md_file.write(text)
            output_suffix = '.docx' if output_format == 'docx' else '.pdf'
            with tempfile.NamedTemporaryFile(suffix=output_suffix, delete=False) as out_file:
                out_path = out_file.name
            cmd = ['pandoc', md_path, '-o', out_path, '--standalone']
            if output_format == 'pdf':
                cmd.extend(['--pdf-engine', engine])
            cmd.extend(extra_args)
            subprocess.run(cmd, check=True, capture_output=True, text=True)
            with open(out_path, 'rb') as f:
                result_bytes = f.read()
            os.unlink(md_path)
            os.unlink(out_path)
            return result_bytes
        except subprocess.CalledProcessError as e:
            continue
        except Exception as e:
            continue
    return None

def create_download_link(text, filename, format):
    if format == "txt":
        b = BytesIO()
        b.write(text.encode('utf-8'))
        b.seek(0)
        return b, f"{filename}.txt", "text/plain"

    elif format == "docx":
        docx_bytes = pandoc_convert(text, 'docx', extra_args=[])
        if docx_bytes:
            b = BytesIO(docx_bytes)
            b.seek(0)
            return b, f"{filename}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = Document()
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(12)
            else:
                doc.add_paragraph()
        b = BytesIO()
        doc.save(b)
        b.seek(0)
        return b, f"{filename}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif format == "pdf":
        pdf_bytes = pandoc_convert(text, 'pdf', extra_args=[])
        if pdf_bytes:
            b = BytesIO(pdf_bytes)
            b.seek(0)
            return b, f"{filename}.pdf", "application/pdf"
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        style_normal = styles['Normal']
        style_normal.fontName = 'Helvetica'
        style_normal.fontSize = 12
        style_normal.leading = 14
        style_normal.alignment = TA_LEFT
        story = []
        for line in text.split('\n'):
            if line.strip() == '':
                story.append(Spacer(1, 12))
            else:
                story.append(Paragraph(line, style_normal))
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        b = BytesIO(pdf_bytes)
        b.seek(0)
        return b, f"{filename}.pdf", "application/pdf"

    else:
        raise ValueError("Formato não suportado")